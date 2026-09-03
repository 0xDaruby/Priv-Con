"""Replaceable local engines for PDF-to-Word conversion."""

from __future__ import annotations

import re
from collections import defaultdict
from collections.abc import Callable, Iterator
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Protocol
from urllib.parse import unquote, urlsplit

import pymupdf
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from pdf2docx import Converter

CancellationCheck = Callable[[], None]
EDITABLE_MODE = "editable"
PRESERVE_APPEARANCE_MODE = "preserve_appearance"
SUPPORTED_PDF_TO_WORD_MODES = frozenset({EDITABLE_MODE, PRESERVE_APPEARANCE_MODE})
_WORD_TOKEN = re.compile(r"[\w@.+-]+", re.UNICODE)
_HTTP_SCHEMES = {"http", "https", "mailto"}
_APPEARANCE_DPI = 144


class PdfToWordEngine(Protocol):
    """Conversion-engine boundary kept intentionally small and replaceable."""

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        *,
        cancellation_check: CancellationCheck | None = None,
    ) -> None: ...


class Pdf2DocxLayoutEngine:
    """Reconstruct editable Word layout with the local pdf2docx engine."""

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        *,
        cancellation_check: CancellationCheck | None = None,
    ) -> None:
        if cancellation_check is not None:
            cancellation_check()

        converter = Converter(str(input_path))

        try:
            # pdf2docx defaults to ignoring individual page errors. PrivCon must
            # fail clearly instead of returning a silently incomplete document.
            converter.convert(
                str(output_path),
                ignore_page_error=False,
                multi_processing=False,
            )
        finally:
            converter.close()

        if cancellation_check is not None:
            cancellation_check()

        _restore_pdf_hyperlinks(input_path, output_path)


class PreserveAppearanceEngine:
    """Place a high-resolution rendering of each PDF page into a DOCX page."""

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        *,
        cancellation_check: CancellationCheck | None = None,
    ) -> None:
        document = Document()
        pdf = pymupdf.open(input_path)

        try:
            for page_index, page in enumerate(pdf):
                if cancellation_check is not None:
                    cancellation_check()

                if page_index == 0:
                    section = document.sections[0]
                    paragraph = document.add_paragraph()
                else:
                    section = document.add_section(WD_SECTION.NEW_PAGE)
                    paragraph = document.add_paragraph()

                page_width = float(page.rect.width)
                page_height = float(page.rect.height)
                _configure_full_page_section(section, page_width, page_height)
                _configure_image_paragraph(paragraph)

                scale = _APPEARANCE_DPI / 72
                pixmap = page.get_pixmap(
                    matrix=pymupdf.Matrix(scale, scale),
                    alpha=False,
                )
                image_stream = BytesIO(pixmap.tobytes("png"))
                run = paragraph.add_run()
                image_scale = min(
                    max(1, page_width - 4) / page_width,
                    max(1, page_height - 4) / page_height,
                )
                run.add_picture(
                    image_stream,
                    width=Pt(page_width * image_scale),
                    height=Pt(page_height * image_scale),
                )

            if cancellation_check is not None:
                cancellation_check()

            document.save(output_path)
        finally:
            pdf.close()


def get_pdf_to_word_engine(mode: str) -> PdfToWordEngine:
    """Return the configured local engine without leaking it into API routes."""
    if mode == EDITABLE_MODE:
        return Pdf2DocxLayoutEngine()
    if mode == PRESERVE_APPEARANCE_MODE:
        return PreserveAppearanceEngine()
    raise ValueError(f"Unsupported PDF-to-Word mode: {mode}")


def _configure_full_page_section(
    section,
    page_width: float,
    page_height: float,
) -> None:
    section.page_width = Pt(page_width)
    section.page_height = Pt(page_height)
    section.top_margin = Pt(0)
    section.right_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def _configure_image_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    # A fixed point-size line height clips inline pictures in LibreOffice.
    # Single spacing expands the line box to the full image height.
    paragraph.paragraph_format.line_spacing = 1.0


def _restore_pdf_hyperlinks(input_path: Path, output_path: Path) -> None:
    """Restore external link annotations omitted by pdf2docx 0.5.8.

    The engine often leaves an empty run where linked text appeared. Prefer
    replacing that placeholder; otherwise use surrounding source-line text to
    attach the link to the closest reconstructed paragraph.
    """
    links = _collect_pdf_links(input_path)

    if not links:
        return

    document = Document(output_path)
    _repair_nested_hyperlinks(document)
    paragraphs = list(_iter_document_paragraphs(document))
    existing_targets = {
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype == RT.HYPERLINK
    }

    for link in links:
        uri = link["uri"]
        label = link["label"]

        if uri in existing_targets:
            continue

        paragraph = _find_link_paragraph(
            paragraphs,
            link["contexts"],
            label,
        )

        if paragraph is None:
            paragraph = document.add_paragraph()

        if _wrap_existing_link_text(paragraph, label, uri):
            existing_targets.add(uri)
            continue

        empty_run = next((run for run in paragraph.runs if not run.text), None)

        if empty_run is not None:
            _replace_empty_run_with_hyperlink(empty_run, label, uri)
        else:
            if paragraph.text.strip():
                paragraph.add_run("  |  ")
            _append_hyperlink(paragraph, label, uri)

        existing_targets.add(uri)

    document.save(output_path)


def _repair_nested_hyperlinks(document) -> None:
    """Move pdf2docx hyperlinks out of invalid parent run elements."""
    nested_hyperlinks = list(document.element.xpath(".//w:r/w:hyperlink"))

    for hyperlink in nested_hyperlinks:
        parent_run = hyperlink.getparent()

        if parent_run is None or parent_run.tag != qn("w:r"):
            continue

        parent_properties = parent_run.find(qn("w:rPr"))

        for child_run in hyperlink.findall(qn("w:r")):
            child_properties = child_run.find(qn("w:rPr"))

            if child_properties is not None:
                child_run.remove(child_properties)
            if parent_properties is not None:
                child_run.insert(0, deepcopy(parent_properties))

        parent_run.addnext(hyperlink)
        parent_run.getparent().remove(parent_run)


def _collect_pdf_links(input_path: Path) -> list[dict[str, object]]:
    grouped: dict[str, dict[str, object]] = {}
    pdf = pymupdf.open(input_path)

    try:
        for page_index, page in enumerate(pdf):
            words = page.get_text("words", sort=True)

            for annotation in page.get_links():
                uri = str(annotation.get("uri") or "").strip()
                parsed = urlsplit(uri)

                if not uri or parsed.scheme.lower() not in _HTTP_SCHEMES:
                    continue

                rectangle = pymupdf.Rect(annotation["from"])
                context = _line_context_for_rectangle(words, rectangle)
                item = grouped.setdefault(
                    uri,
                    {
                        "uri": uri,
                        "label": _link_label(uri),
                        "contexts": [],
                        "order": (page_index, rectangle.y0, rectangle.x0),
                    },
                )
                contexts = item["contexts"]
                assert isinstance(contexts, list)
                if context:
                    contexts.append(context)

                order = item["order"]
                if isinstance(order, tuple):
                    item["order"] = min(
                        order,
                        (page_index, rectangle.y0, rectangle.x0),
                    )
    finally:
        pdf.close()

    return sorted(grouped.values(), key=lambda item: item["order"])


def _line_context_for_rectangle(words: list[tuple], rectangle) -> str:
    matching_lines: dict[tuple[int, int], list[tuple[int, str]]] = defaultdict(list)

    for word in words:
        _x0, y0, _x1, y1, text, block_index, line_index, word_index = word[:8]
        word_midpoint = (float(y0) + float(y1)) / 2

        if rectangle.y0 - 1 <= word_midpoint <= rectangle.y1 + 1:
            matching_lines[(int(block_index), int(line_index))].append(
                (int(word_index), str(text))
            )

    lines = [
        " ".join(text for _, text in sorted(line_words))
        for line_words in matching_lines.values()
    ]
    return " ".join(lines)


def _link_label(uri: str) -> str:
    parsed = urlsplit(uri)

    if parsed.scheme.lower() == "mailto":
        return unquote(parsed.path)

    host = parsed.netloc.removeprefix("www.")
    path = unquote(parsed.path).rstrip("/")
    return f"{host}{path}" or uri


def _find_link_paragraph(
    paragraphs: list[Paragraph],
    contexts: object,
    label: str,
) -> Paragraph | None:
    context_values = contexts if isinstance(contexts, list) else []
    label_tokens = set(_normalized_tokens(label))
    best: tuple[float, Paragraph] | None = None

    for paragraph in paragraphs:
        paragraph_tokens = set(_normalized_tokens(paragraph.text))

        if not paragraph_tokens:
            continue

        for context in context_values:
            context_tokens = set(_normalized_tokens(str(context))) - label_tokens

            if not context_tokens:
                continue

            score = len(paragraph_tokens & context_tokens) / len(context_tokens)

            if best is None or score > best[0]:
                best = (score, paragraph)

    return best[1] if best is not None and best[0] >= 0.2 else None


def _normalized_tokens(value: str) -> list[str]:
    return [token.casefold() for token in _WORD_TOKEN.findall(value)]


def _iter_document_paragraphs(document) -> Iterator[Paragraph]:
    yield from document.paragraphs

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _wrap_existing_link_text(
    paragraph: Paragraph,
    label: str,
    uri: str,
) -> bool:
    folded_label = label.casefold()

    for run in paragraph.runs:
        start = run.text.casefold().find(folded_label)

        if start < 0:
            continue

        before = run.text[:start]
        display = run.text[start : start + len(label)]
        after = run.text[start + len(label) :]
        template_properties = deepcopy(run._r.rPr) if run._r.rPr is not None else None
        run.text = before
        hyperlink = _hyperlink_element(paragraph, display, uri, template_properties)
        run._r.addnext(hyperlink)

        if after:
            after_run = OxmlElement("w:r")
            if template_properties is not None:
                after_run.append(deepcopy(template_properties))
            text = OxmlElement("w:t")
            text.set(qn("xml:space"), "preserve")
            text.text = after
            after_run.append(text)
            hyperlink.addnext(after_run)

        return True

    return False


def _replace_empty_run_with_hyperlink(run: Run, label: str, uri: str) -> None:
    paragraph = run._parent
    template_properties = deepcopy(run._r.rPr) if run._r.rPr is not None else None
    hyperlink = _hyperlink_element(paragraph, label, uri, template_properties)
    run._r.addnext(hyperlink)
    run._r.getparent().remove(run._r)


def _append_hyperlink(paragraph: Paragraph, label: str, uri: str) -> None:
    template_run = next((run for run in paragraph.runs if run.text), None)
    template_properties = (
        deepcopy(template_run._r.rPr)
        if template_run is not None and template_run._r.rPr is not None
        else None
    )
    paragraph._p.append(_hyperlink_element(paragraph, label, uri, template_properties))


def _hyperlink_element(
    paragraph: Paragraph,
    label: str,
    uri: str,
    run_properties,
):
    relationship_id = paragraph.part.relate_to(
        uri,
        RT.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")

    if run_properties is not None:
        run.append(deepcopy(run_properties))

    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    return hyperlink
