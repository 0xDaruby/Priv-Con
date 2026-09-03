"""Local layout-aware PDF-to-Word conversion tests."""

import zipfile
from io import BytesIO
from pathlib import Path

import pymupdf
import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

from app.config import settings
from app.core.exceptions import ConversionError, ValidationError
from app.services import pdf_to_word_engines, pdf_to_word_service
from app.services.pdf_to_word_engines import (
    EDITABLE_MODE,
    PRESERVE_APPEARANCE_MODE,
)


@pytest.fixture()
def conversion_paths(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    output_dir = tmp_path / "outputs"
    output_dir.mkdir()
    input_path = tmp_path / "source.pdf"
    _make_layout_fixture(input_path)
    monkeypatch.setattr(settings, "output_temp_dir", output_dir)
    return input_path, output_dir


def test_pdf_to_docx_uses_editable_layout_engine_by_default(
    conversion_paths,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    input_path, _ = conversion_paths
    calls: list[tuple[str, Path]] = []

    class FakeLayoutEngine:
        def convert(
            self,
            _input_path: Path,
            output_path: Path,
            *,
            cancellation_check=None,
        ) -> None:
            calls.append((EDITABLE_MODE, output_path))
            if cancellation_check is not None:
                cancellation_check()
            document = Document()
            document.add_paragraph("Layout-aware content")
            document.save(output_path)

    monkeypatch.setattr(
        pdf_to_word_service,
        "get_pdf_to_word_engine",
        lambda mode: FakeLayoutEngine(),
    )

    output_path = pdf_to_word_service.pdf_to_docx(input_path, "job-1")

    assert calls == [(EDITABLE_MODE, output_path)]
    assert "Layout-aware content" in Document(output_path).paragraphs[0].text


def test_pdf_to_docx_rejects_unknown_mode_without_creating_output(
    conversion_paths,
) -> None:
    input_path, output_dir = conversion_paths

    with pytest.raises(ValidationError, match="Choose Editable Word"):
        pdf_to_word_service.pdf_to_docx(
            input_path,
            "job-invalid",
            mode="plain_text",
        )

    assert list(output_dir.iterdir()) == []


def test_editable_engine_failure_does_not_fall_back_to_plain_text(
    conversion_paths,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    input_path, output_dir = conversion_paths

    class FailingLayoutEngine:
        def convert(self, *args, **kwargs) -> None:
            raise RuntimeError("layout parser stopped")

    monkeypatch.setattr(
        pdf_to_word_service,
        "get_pdf_to_word_engine",
        lambda mode: FailingLayoutEngine(),
    )

    with pytest.raises(ConversionError) as raised:
        pdf_to_word_service.pdf_to_docx(input_path, "job-failure")

    assert raised.value.code == "layout_conversion_failed"
    assert "Preserve Appearance" in raised.value.message
    assert list(output_dir.iterdir()) == []


def test_preserve_appearance_creates_one_rendered_image_per_pdf_page(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    input_path = tmp_path / "two-pages.pdf"
    pdf = pymupdf.open()
    first = pdf.new_page(width=420, height=600)
    first.insert_text((40, 80), "First visual page")
    second = pdf.new_page(width=600, height=420)
    second.draw_rect(pymupdf.Rect(50, 50, 300, 200), fill=(0.1, 0.7, 0.3))
    pdf.save(input_path)
    pdf.close()
    output_dir = tmp_path / "outputs"
    output_dir.mkdir()
    monkeypatch.setattr(settings, "output_temp_dir", output_dir)

    output_path = pdf_to_word_service.pdf_to_docx(
        input_path,
        "job-appearance",
        mode=PRESERVE_APPEARANCE_MODE,
    )
    document = Document(output_path)
    image_relationships = [
        relationship
        for relationship in document.part.rels.values()
        if relationship.reltype == RT.IMAGE
    ]
    assert len(document.sections) == 2
    assert len(image_relationships) == 2


def test_editable_engine_preserves_layout_content_and_hyperlink(
    conversion_paths,
) -> None:
    input_path, _ = conversion_paths

    output_path = pdf_to_word_service.pdf_to_docx(
        input_path,
        "job-layout",
        mode=EDITABLE_MODE,
    )
    document = Document(output_path)
    document_text = "\n".join(
        paragraph.text
        for paragraph in pdf_to_word_engines._iter_document_paragraphs(document)
    )
    hyperlink_targets = {
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype == RT.HYPERLINK
    }
    image_relationships = [
        relationship
        for relationship in document.part.rels.values()
        if relationship.reltype == RT.IMAGE
    ]
    with zipfile.ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")

    assert "Left column uses Courier" in document_text
    assert "Right column uses Times Italic" in document_text
    assert "Table A1" in document_text
    assert "privcon.local/docs" in document_xml
    assert "https://privcon.local/docs" in hyperlink_targets
    assert not document.element.xpath(".//w:r/w:hyperlink")
    assert image_relationships
    assert document.tables


def _make_layout_fixture(path: Path) -> None:
    pdf = pymupdf.open()
    page = pdf.new_page(width=612, height=792)
    page.draw_rect(
        pymupdf.Rect(36, 36, 576, 82),
        color=(0.05, 0.25, 0.15),
        fill=(0.85, 0.96, 0.9),
        width=2,
    )
    page.insert_text((50, 65), "LAYOUT FIDELITY FIXTURE", fontsize=18)
    page.insert_textbox(
        pymupdf.Rect(48, 110, 285, 205),
        "Left column uses Courier\nwith independent positioning.",
        fontname="cour",
        fontsize=11,
    )
    page.insert_textbox(
        pymupdf.Rect(327, 110, 564, 205),
        "Right column uses Times Italic\nwith a second text flow.",
        fontname="Times-Italic",
        fontsize=11,
    )

    table_rect = pymupdf.Rect(48, 245, 564, 350)
    page.draw_rect(table_rect, color=(0, 0, 0), width=1)
    page.draw_line((306, 245), (306, 350), color=(0, 0, 0), width=1)
    page.draw_line((48, 297.5), (564, 297.5), color=(0, 0, 0), width=1)
    page.insert_text((60, 278), "Table A1")
    page.insert_text((320, 278), "Table B1")
    page.insert_text((60, 330), "Table A2")
    page.insert_text((320, 330), "Table B2")

    image_bytes = BytesIO()
    image = Image.new("RGB", (100, 60), (30, 130, 90))
    image.save(image_bytes, format="PNG")
    image.close()
    page.insert_image(
        pymupdf.Rect(48, 390, 198, 480),
        stream=image_bytes.getvalue(),
    )

    link_label = "privcon.local/docs"
    page.insert_text((48, 535), link_label, color=(0, 0, 0.8))
    page.insert_link(
        {
            "kind": pymupdf.LINK_URI,
            "from": pymupdf.Rect(48, 520, 160, 540),
            "uri": "https://privcon.local/docs",
        }
    )
    pdf.save(path)
    pdf.close()
