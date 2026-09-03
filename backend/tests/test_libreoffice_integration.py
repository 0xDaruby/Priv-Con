"""Optional end-to-end smoke test against the installed LibreOffice binary."""

from pathlib import Path

import pytest
from docx import Document

from app.config import settings
from app.core.exceptions import ConversionError
from app.core.file_utils import new_job_id
from app.core.validators import validate_office_structure, validate_pdf_structure
from app.services import libreoffice_service
from app.services.cleanup_service import cleanup_now
from app.services.libreoffice_service import (
    _resolve_libreoffice_executable,
    convert_to_pdf,
)
from app.services.pdf_to_word_service import pdf_to_docx


def test_installed_libreoffice_converts_a_real_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    try:
        executable = _resolve_libreoffice_executable()
    except ConversionError:
        pytest.skip("LibreOffice is not installed in this test environment")

    upload_dir = tmp_path / "uploads"
    output_dir = tmp_path / "outputs"
    upload_dir.mkdir()
    output_dir.mkdir()
    monkeypatch.setattr(settings, "upload_temp_dir", upload_dir)
    monkeypatch.setattr(settings, "output_temp_dir", output_dir)
    monkeypatch.setattr(settings, "libreoffice_path", executable)
    input_path = upload_dir / "smoke-test.docx"
    document = Document()
    document.add_heading("PrivCon", level=1)
    document.add_paragraph("Local conversion security smoke test.")
    document.save(input_path)
    output_path: Path | None = None
    word_output_path: Path | None = None
    process_diagnostics: list[tuple[int, str, str]] = []
    real_run = libreoffice_service.subprocess.run

    def capture_conversion_result(*args, **kwargs):
        result = real_run(*args, **kwargs)
        process_diagnostics.append((result.returncode, result.stdout, result.stderr))
        return result

    monkeypatch.setattr(
        libreoffice_service.subprocess,
        "run",
        capture_conversion_result,
    )

    try:
        validate_office_structure(input_path, "docx")
        try:
            output_path = convert_to_pdf(input_path, new_job_id())
        except ConversionError as exc:
            raise AssertionError(
                f"LibreOffice conversion failed: {process_diagnostics!r}"
            ) from exc
        assert validate_pdf_structure(output_path) == 1
        word_output_path = pdf_to_docx(output_path, new_job_id())
        validate_office_structure(word_output_path, "docx")
        round_trip_text = "\n".join(
            paragraph.text for paragraph in Document(word_output_path).paragraphs
        )
        assert "PrivCon" in round_trip_text
        assert "Local conversion security smoke test." in round_trip_text
    finally:
        cleanup_targets = [input_path]

        if output_path is not None:
            cleanup_targets.append(output_path.parent)

        if word_output_path is not None:
            cleanup_targets.append(word_output_path.parent)

        cleanup_now(cleanup_targets)

    assert list(upload_dir.iterdir()) == []
    assert list(output_dir.iterdir()) == []
